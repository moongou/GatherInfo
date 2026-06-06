"""
Report Export — render a generated Report's markdown into MD / HTML / DOCX / PDF
files on disk, organized into a date-based subfolder.

Design notes:
    - CJK support: PDF uses reportlab's built-in CID font "STSong-Light" (no external
      font files needed). DOCX uses python-docx defaults which already handle CJK.
    - Heavy deps (reportlab, python-docx) are imported lazily so a missing optional
      dependency only disables that one format instead of breaking the whole app.
    - Output layout: {root}/{date_pattern}/{safe_title}.{ext}
      root default = <repo>/data/reports
"""
from __future__ import annotations

import html as _html
import os
import re
from datetime import datetime

from app.database import DATA_DIR
from app.models import Report, SystemConfig, Topic

SUPPORTED_FORMATS = ("md", "html", "docx", "pdf")


# ── Public API ──────────────────────────────────────────────────────────────

def export_report(report: Report, system: SystemConfig | None, topic: Topic | None) -> dict[str, str]:
    """Render `report` to all configured formats. Returns {format: abs_path}.

    Skips formats whose optional dependency is unavailable (records nothing for them).
    """
    formats = _resolve_formats(system)
    title = _resolve_title(report, system, topic)
    safe_title = _safe_filename(title)

    root = (system.report_output_dir if system and system.report_output_dir else None) \
        or os.path.join(DATA_DIR, "reports")
    pattern = (system.report_dir_pattern if system and system.report_dir_pattern else "%Y-%m-%d")
    subdir = datetime.now().strftime(pattern)
    out_dir = os.path.join(root, subdir)
    os.makedirs(out_dir, exist_ok=True)

    body = report.content or ""
    out_files: dict[str, str] = {}

    for fmt in formats:
        path = os.path.join(out_dir, f"{safe_title}.{fmt}")
        try:
            if fmt == "md":
                _write_md(path, title, body)
            elif fmt == "html":
                _write_html(path, title, body)
            elif fmt == "docx":
                _write_docx(path, title, body)
            elif fmt == "pdf":
                _write_pdf(path, title, body)
            else:
                continue
            out_files[fmt] = path
        except ImportError:
            # Optional dependency missing — skip this format silently.
            continue
        except Exception:
            # Never let one format failure abort the others.
            continue

    report.output_files = out_files
    report.output_dir = out_dir
    return out_files


# ── Helpers ─────────────────────────────────────────────────────────────────

def _resolve_formats(system: SystemConfig | None) -> list[str]:
    fmts = (system.report_formats if system and system.report_formats else None) or list(SUPPORTED_FORMATS)
    return [f for f in fmts if f in SUPPORTED_FORMATS]


def _resolve_title(report: Report, system: SystemConfig | None, topic: Topic | None) -> str:
    fmt = (system.report_title_format if system and system.report_title_format else None) \
        or "{topic}_情报报告_{date}"
    topic_name = (topic.name if topic else None) or report.topic_id or "报告"
    date_str = datetime.now().strftime("%Y-%m-%d")
    try:
        return fmt.format(topic=topic_name, date=date_str, title=report.title or topic_name)
    except (KeyError, IndexError, ValueError):
        return report.title or f"{topic_name}_{date_str}"


_UNSAFE = re.compile(r'[\\/:*?"<>|\r\n\t]+')


def _safe_filename(name: str) -> str:
    cleaned = _UNSAFE.sub("_", name).strip(" .")
    cleaned = cleaned[:120] if len(cleaned) > 120 else cleaned
    return cleaned or "report"


def _write_md(path: str, title: str, body: str) -> None:
    text = body if body.lstrip().startswith("#") else f"# {title}\n\n{body}"
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(text)


def _write_html(path: str, title: str, body: str) -> None:
    content_html = _markdown_to_html(body)
    doc = (
        "<!DOCTYPE html>\n<html lang=\"zh-CN\">\n<head>\n<meta charset=\"utf-8\">\n"
        f"<title>{_html.escape(title)}</title>\n"
        "<style>\n"
        "body{font-family:-apple-system,'Segoe UI','PingFang SC','Microsoft YaHei',sans-serif;"
        "max-width:860px;margin:40px auto;padding:0 20px;line-height:1.7;color:#1f2937;}\n"
        "h1,h2,h3{color:#111827;margin-top:1.4em;}\n"
        "code{background:#f3f4f6;padding:2px 5px;border-radius:4px;}\n"
        "pre{background:#f3f4f6;padding:12px;border-radius:8px;overflow:auto;}\n"
        "blockquote{border-left:3px solid #d1d5db;margin:0;padding-left:14px;color:#4b5563;}\n"
        "table{border-collapse:collapse;}td,th{border:1px solid #d1d5db;padding:6px 10px;}\n"
        "</style>\n</head>\n<body>\n"
        f"{content_html}\n</body>\n</html>\n"
    )
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(doc)


def _write_docx(path: str, title: str, body: str) -> None:
    from docx import Document  # lazy import

    doc = Document()
    doc.add_heading(title, level=0)
    for line in body.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
        else:
            doc.add_paragraph(_strip_inline_md(stripped))
    doc.save(path)


def _write_pdf(path: str, title: str, body: str) -> None:
    from reportlab.lib.pagesizes import A4  # lazy import
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    font_name = "STSong-Light"
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    except Exception:
        font_name = "Helvetica"

    base = getSampleStyleSheet()
    body_style = ParagraphStyle("Body", parent=base["BodyText"], fontName=font_name,
                                fontSize=10.5, leading=16)
    h1 = ParagraphStyle("H1", parent=base["Heading1"], fontName=font_name, fontSize=18, leading=24)
    h2 = ParagraphStyle("H2", parent=base["Heading2"], fontName=font_name, fontSize=14, leading=20)
    h3 = ParagraphStyle("H3", parent=base["Heading3"], fontName=font_name, fontSize=12, leading=18)
    title_style = ParagraphStyle("Title", parent=base["Title"], fontName=font_name, fontSize=22, leading=28)

    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=20 * mm, rightMargin=20 * mm,
                            topMargin=18 * mm, bottomMargin=18 * mm)
    flow = [Paragraph(_html.escape(title), title_style), Spacer(1, 8)]
    for line in body.splitlines():
        stripped = line.strip()
        if not stripped:
            flow.append(Spacer(1, 6))
            continue
        if stripped.startswith("### "):
            flow.append(Paragraph(_html.escape(stripped[4:]), h3))
        elif stripped.startswith("## "):
            flow.append(Paragraph(_html.escape(stripped[3:]), h2))
        elif stripped.startswith("# "):
            flow.append(Paragraph(_html.escape(stripped[2:]), h1))
        elif stripped.startswith(("- ", "* ")):
            flow.append(Paragraph("• " + _html.escape(_strip_inline_md(stripped[2:])), body_style))
        else:
            flow.append(Paragraph(_html.escape(_strip_inline_md(stripped)), body_style))
    doc.build(flow)


# ── Minimal markdown helpers ─────────────────────────────────────────────────

def _strip_inline_md(text: str) -> str:
    """Remove the most common inline markdown markers for plain-text renderers."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1 (\2)", text)
    return text


def _markdown_to_html(body: str) -> str:
    """Very small markdown→HTML for headings, lists, bold/italic/code and paragraphs."""
    lines = body.splitlines()
    out: list[str] = []
    in_list = False

    def close_list() -> None:
        nonlocal in_list
        if in_list:
            out.append("</ul>")
            in_list = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            close_list()
            continue
        if stripped.startswith("### "):
            close_list(); out.append(f"<h3>{_inline_html(stripped[4:])}</h3>")
        elif stripped.startswith("## "):
            close_list(); out.append(f"<h2>{_inline_html(stripped[3:])}</h2>")
        elif stripped.startswith("# "):
            close_list(); out.append(f"<h1>{_inline_html(stripped[2:])}</h1>")
        elif stripped.startswith(("- ", "* ")):
            if not in_list:
                out.append("<ul>"); in_list = True
            out.append(f"<li>{_inline_html(stripped[2:])}</li>")
        else:
            close_list(); out.append(f"<p>{_inline_html(stripped)}</p>")
    close_list()
    return "\n".join(out)


def _inline_html(text: str) -> str:
    text = _html.escape(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
    text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r'<a href="\2">\1</a>', text)
    return text
